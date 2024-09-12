import streamlit as st
from docx import Document
import google.generativeai as genai
import os
import asyncio
import io
import json
import typing_extensions as typing

# Add custom CSS to hide the GitHub icon
hide_streamlit_style = """
<style>
#MainMenu {visibility: hidden;}
footer {visibility: hidden;}
header {visibility: hidden;}
.main {margin-top: 0px;}
</style>

"""
st.markdown(hide_streamlit_style, unsafe_allow_html=True) 

# --- 全局变量 ---
# 从 Vercel 环境变量中获取 API 密钥
api_key = os.environ.get("GOOGLE_API_KEY")

# 检查是否成功获取 API 密钥
if not api_key:
    st.error("GOOGLE_API_KEY environment variable is not set!")
    st.stop()  # 停止应用加载

class translatePair(typing.TypedDict):
    index: str
    paragraph_index: str
    translation: str

# --- 全局变量 ---

# 获取用户选择的语言，默认为英文
if 'language' not in st.session_state:
    st.session_state['language'] = 'en'

# --- 语言选择器 ---
# 语言选项
language_options = {
    "zh": "中文",
    "en": "English",
    "ja": "日本語"
}

# 创建语言选择下拉菜单
selected_language = st.selectbox(
    "Select Language / 语言选择 / 言語を選択",
    list(language_options.keys()),
    format_func=lambda x: language_options[x],
    key='language'  # 将用户的选择存储到 st.session_state['language']
)

# 从 session state 获取语言 (如果用户已经选择了语言，则使用用户的选择)
if 'language' in st.session_state:
    selected_language = st.session_state['language']

def load_locale(language='en'):
    """加载指定语言的文本"""
    try:
        with open(f'locales/{language}.json', 'r', encoding='utf-8') as f:
            return json.load(f)
    except FileNotFoundError:
        return {}  # 返回空字典，如果找不到语言文件
    
# 加载语言文件
locale = load_locale(selected_language)

# Streamlit 应用标题
st.title(locale.get("title", "Word Document Translation"))

# 文件上传
uploaded_file = st.file_uploader(locale.get("upload_file", "Upload Word Document (.docx)"), type=["docx","doc"])

# 目标语言选择
target_language = st.selectbox(
    locale.get("select_target_language", "Select Target Language"), ["zh-CN", "en", "ja", "ko", "fr", "de", "es"]
)

# 双语模式选择
bilingual = st.checkbox(locale.get("bilingual_mode", "Bilingual Mode"), False)

# --- 翻译进度条 ---
progress_bar = st.progress(0, text=locale.get("preparing", "Preparing..."))
progress_lock = asyncio.Lock()

# 设置 API 密钥
genai.configure(api_key=api_key)

# 选择 Gemini Pro 模型
model = genai.GenerativeModel(
    "gemini-1.5-flash",
    generation_config={
        "response_mime_type": "application/json",
        "response_schema": list[translatePair],
    },
)

async def update_progress(completed):
    """更新进度条"""
    async with progress_lock:
        progress = int(completed)
        progress_bar.progress(progress, text=locale.get("completed", "Completed ({}%)...").format(progress))

async def update_progress_text(completed,text):
    """更新进度条"""
    async with progress_lock:
        progress = int(completed)
        progress_bar.progress(progress, text)


async def translate_text(texts, start_progress, end_progress):
    """使用 Google Gemini API 批量翻译文本"""
    # 设置最大 token 限制
    max_tokens = 10000

    # 构建通用的 prompt
    prompt = f"""
                Translate the following texts to {target_language}, paying close attention to the context and ensuring accuracy. Double-check for any potentially ambiguous words or phrases and choose the most appropriate translation. 


                **Examples of potential ambiguities:**
                - If the word "charge" refers to billing, ensure it is not translated as "charging" (as in electricity).

                **Formatting instructions:**
                - Do not add any extra line breaks, markdown formatting, numbering, or any other special formatting. 
                - Please preserving all original formatting, including spaces, line breaks, and special characters such as tabs.
                - Directly return a JSON array without any additional formatting. 
                - The returned JSON array must strictly adhere to the following JSON format, each object in array must include index, paragraph_index and translation.  It is absolutely forbidden to return only the translated text directly.
                - Make sure the output is a complete and valid JSON array.

                Only return the result in the following JSON format,replace translation value with the translated text :
                [
                  {{"index": "0", "paragraph_index":"0", "translation": "Translated text 1"}},
                  {{"index": "1", "paragraph_index":"1", "translation": "Translated text 2"}}
                ]

                Please translate the following texts:
                """
    # 将文本列表分割成多个批次，每个批次的 token 数量不超过最大限制
    batches = []
    current_batch = []
    current_tokens = 0
    for i, (paragraph_index, text) in enumerate(texts):  # 获取段落索引
        text_tokens = len(text)
        if current_tokens + text_tokens > max_tokens:
            batches.append(current_batch)
            current_batch = []
            current_tokens = 0
        current_batch.append((i, paragraph_index, text))  # 存储段落索引、批次索引和文本
        current_tokens += text_tokens
    if current_batch:
        batches.append(current_batch)

    # 依次翻译每个批次
    translations = []
    total_batches = len(batches)
    max_retries = 5  # 设置最大重试次数
    # 创建一个空的 Streamlit 组件，用于存储错误信息
    error_message = st.empty() 
    for batch_index, batch in enumerate(batches):
        retry_count = 0
        # 计算当前批次的进度范围
        batch_start = int(batch_index / total_batches * 100)
        batch_end = int((batch_index + 1) / total_batches * 100)

        # 计算实际进度范围
        actual_start = start_progress + batch_start * (end_progress - start_progress) / 100
        actual_end = start_progress + batch_end * (end_progress - start_progress) / 100

        batch_size = actual_end - actual_start
        await update_progress(actual_start + 0.2 * batch_size)

        batch_prompt = prompt  # 使用通用的 prompt
        for i, paragraph_index, text in batch:  # 获取段落索引
            batch_prompt += f'{{"index": {i}, "paragraph_index": {paragraph_index}, "translation": "{text}"}}\n'  # 添加段落索引到 JSON
        
        while retry_count < max_retries:
            try:
                response = model.generate_content(batch_prompt)
                batch_translations = response.text
                await update_progress(actual_start + 0.9 * batch_size)

                # 移除 ```json ``` 包装
                if batch_translations.startswith("```json\n") and batch_translations.endswith("\n```"):
                    batch_translations = batch_translations[8:-4].strip()

                # 使用 replace() 方法替换无效字符
                batch_translations = ''.join(
                    c for c in batch_translations if c.isprintable()
                )

                batch_translations = json.loads(batch_translations)
                translations.extend(batch_translations)
                await update_progress(actual_end)
                error_message.empty()  # 清除错误信息
                break  # 翻译成功，退出循环
            except Exception as e:
                retry_count += 1
                st.warning(f"Error parsing JSON for batch {batch_index + 1} , retrying attempt {retry_count} ...")  # 显示完整的错误堆栈
                
        if retry_count == max_retries:
            error_message.error(
                f"Batch {batch_index + 1} translation failed, maximum retries ({max_retries}) reached"
            )
            raise Exception(f"Batch {batch_index + 1} translation failed")  # 抛出异常

    translations.sort(key=lambda x: int(x["index"]))  # 将 index 转换为整数
    return translations


async def process_paragraph(paragraph, translations, paragraph_index):
    """处理单个段落，包含翻译和进度更新"""
    if paragraph.runs:  # 检查段落中是否有 Run
        # 找到第一个有文本内容的 Run
        for run in paragraph.runs:
            if run:
                break
        else:
            # 如果没有找到有文本内容的 Run，则跳过该段落
            return

        if bilingual:
            # 在 translations 中查找对应的翻译结果
            translated_text = next(
                (t["translation"].rstrip() for t in translations if int(t["paragraph_index"]) == paragraph_index), 
                None,
            )
            if translated_text:
                new_run = paragraph.add_run("\n" + translated_text)
                new_run.font.bold = run.font.bold
                new_run.font.italic = run.font.italic
                new_run.font.underline = run.font.underline
                new_run.font.color.rgb = run.font.color.rgb
        else:            
            try:
                # 在 translations 中查找对应的翻译结果，使用 paragraph_index 作为索引
                translated_text = next(
                    (t["translation"].rstrip() for t in translations if int(t["paragraph_index"]) == paragraph_index),
                    None,
                )
                if translated_text:
                    paragraph.text = translated_text
            except Exception as e:
                st.exception(e)
                raise e  # 抛出异常
            if run:
               for new_run in paragraph.runs:
                   new_run.font.bold = run.font.bold
                   new_run.font.italic = run.font.italic
                   new_run.font.underline = run.font.underline
                   new_run.font.color.rgb = run.font.color.rgb

async def translate_subdocument(document, start_paragraph, end_paragraph, start_progress, end_progress):
    """翻译文档的一部分"""
    global total_elements, completed_elements
    total_elements = end_paragraph - start_paragraph
    texts_to_translate = []

    for i in range(start_paragraph, end_paragraph):
        paragraph = document.paragraphs[i]
        original_text = paragraph.text.strip()
        if original_text:
            texts_to_translate.append((i, paragraph.text))

    completed_elements = 0

    translations = await translate_text(texts_to_translate, start_progress, end_progress)

    for i in range(start_paragraph, end_paragraph):
        paragraph = document.paragraphs[i]
        original_text = paragraph.text.strip()
        if original_text:
            await process_paragraph(paragraph, translations, i)

    return document


async def translate_document(document):
    """将文档分割成多个部分并翻译"""
    update_progress_text(0, text=locale.get("preparing", "Preparing..."))
    # 设置每个部分的最大字节大小
    max_part_size = 1024 * 1024  # 1MB

    # 获取文档内容的字节大小
    doc_content = io.BytesIO()
    document.save(doc_content)
    doc_size = doc_content.tell()

    # 如果文档大小小于 1MB，则不需要分割
    if doc_size < max_part_size:
        return await translate_subdocument(document, 0, len(document.paragraphs), 0, 100)

    # 估算平均段落大小
    total_paragraphs = len(document.paragraphs)
    average_paragraph_size = doc_size / total_paragraphs

    # 计算每个部分的段落数
    paragraphs_per_part = int(max_part_size / average_paragraph_size)

    # 计算需要分割的份数
    num_parts = (total_paragraphs + paragraphs_per_part - 1) // paragraphs_per_part

    current_paragraph = 0
    start_progress = 0
    for i in range(num_parts):
        end_progress = int((i + 1) / num_parts * 100)
        # 计算当前部分的结束段落索引
        part_end_paragraph = min(current_paragraph + paragraphs_per_part, total_paragraphs)
        document = await translate_subdocument(
            document, current_paragraph, part_end_paragraph, start_progress, end_progress
        )
        current_paragraph = part_end_paragraph
        start_progress = end_progress

    update_progress_text(100, text=locale.get("preparing_download", "Preparing for download..."))
    return document

# 初始化翻译状态
is_translating = False

# 翻译按钮
if st.button(locale.get("translate_now", "Translate Now"), disabled=is_translating) and uploaded_file is not None:
    try:
        # 设置翻译状态为 True
        is_translating = True
        
        # 读取 Word 文档
        doc = Document(uploaded_file)

        # 使用 asyncio 运行异步函数
        translated_doc = asyncio.run(translate_document(doc))

        # 将翻译后的文档保存到内存中
        output = io.BytesIO()
        translated_doc.save(output)
        output.seek(0)

        # 提供下载链接
        st.download_button(
            label=locale.get("download", "Download"),
            data=output.getvalue(),
            file_name=f"Translated_{uploaded_file.name}",
        )

    except Exception as e:
        locale.get("error_during_translation", "Error during translation: {}").format(e)
        st.exception(e)  # 显示完整的错误堆栈
        
    finally:
        # 设置翻译状态为 False
        is_translating = False
